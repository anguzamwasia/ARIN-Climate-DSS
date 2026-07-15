from fastapi import APIRouter, Depends, Query, UploadFile, File, HTTPException
from sqlalchemy.orm import Session
from sqlalchemy import func
from typing import Optional
import os
import uuid
from app.database import get_db
from app.schemas.document import DocumentOut
from app.models.document import Document

router = APIRouter()

@router.get("/documents/stats/countries")
def count_unique_countries(db: Session = Depends(get_db)):
    count = db.query(func.count(func.distinct(Document.country)))\
        .filter(Document.source == 'UNFCCC', Document.country != 'Africa (Global)')\
        .scalar() or 0
    return {"count": count}

@router.get("/api/stats")
def get_global_stats(db: Session = Depends(get_db)):
    countries_count = db.query(func.count(func.distinct(Document.country)))\
        .filter(Document.source == 'UNFCCC', Document.country != 'Africa (Global)')\
        .scalar() or 0

    reports_count = db.query(func.count(Document.id))\
        .filter(Document.source.in_(['UNFCCC', 'KNBS', 'World Bank', 'KMD']))\
        .scalar() or 0

    field_count = db.query(func.count(Document.id))\
        .filter(Document.source == 'KOBO')\
        .scalar() or 0

    return {
        "stats": [
            {
                "icon": "globe",
                "value": countries_count,
                "suffix": "+",
                "label": "African Countries",
                "description": "Research coverage"
            },
            {
                "icon": "fileText",
                "value": reports_count,
                "suffix": "+",
                "label": "Policy Reports",
                "description": "UNFCCC & National"
            },
            {
                "icon": "users",
                "value": field_count,
                "suffix": "+",
                "label": "Field Submissions",
                "description": "KoboCollect data"
            }
        ]
    }

@router.get("/api/v1/admin/content/stats")
def get_admin_content_stats(db: Session = Depends(get_db)):
    research_papers = db.query(func.count(Document.id)).filter(Document.source == 'ARIN').scalar() or 0
    media_processed = db.query(func.count(Document.id)).filter(Document.source == 'WHISPER').scalar() or 0
    
    return {
        "research_papers": research_papers,
        "media_processed": media_processed
    }

@router.get("/documents", response_model=list[DocumentOut])
def list_documents(
    source: Optional[str] = Query(None),
    country: Optional[str] = Query(None),
    type: Optional[str] = Query(None),
    limit: int = Query(20, le=10000),
    offset: int = Query(0),
    db: Session = Depends(get_db)
):
    query = db.query(Document)
    if source:
        query = query.filter(Document.source == source)
    if country:
        query = query.filter(Document.country.ilike(f"%{country}%"))
    if type:
        query = query.filter(Document.type.ilike(f"%{type}%"))

    return query.order_by(Document.scraped_at.desc()).offset(offset).limit(limit).all()

@router.get("/documents/{doc_id}", response_model=DocumentOut)
def get_document(doc_id: int, db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc

@router.post("/api/v1/admin/documents/upload")
async def upload_document(file: UploadFile = File(...), db: Session = Depends(get_db)):
    # Define acceptable research paper formats
    allowed_extensions = {".pdf", ".docx", ".csv", ".xlsx"}
    ext = os.path.splitext(file.filename)[1].lower()
    
    if ext not in allowed_extensions:
        raise HTTPException(status_code=400, detail=f"Unsupported file format {ext}. Allowed: pdf, docx, csv, xlsx")

    # Secure storage directory
    upload_dir = "uploads/documents"
    os.makedirs(upload_dir, exist_ok=True)
    
    # Use the original filename instead of UUID
    safe_filename = file.filename
    file_path = os.path.join(upload_dir, safe_filename)

    try:
        content = await file.read()
        with open(file_path, "wb") as f:
            f.write(content)
            
        # Parse text based on extension
        extracted_text = ""
        if ext == ".pdf":
            import fitz
            doc = fitz.open(file_path)
            extracted_text = "\n".join(page.get_text() for page in doc)
        elif ext == ".docx":
            import docx
            doc = docx.Document(file_path)
            extracted_text = "\n".join(para.text for para in doc.paragraphs)
        elif ext == ".csv":
            import pandas as pd
            df = pd.read_csv(file_path)
            extracted_text = df.to_string()
        elif ext == ".xlsx":
            import pandas as pd
            df = pd.read_excel(file_path)
            extracted_text = df.to_string()
            
        from datetime import datetime
        # Insert into Document table
        new_doc = Document(
            title=file.filename,
            source="ARIN",
            country="Africa (Global)", # or some default
            url="#",
            file_url=safe_filename,
            scraped_at=datetime.utcnow(),
            content_text=extracted_text,
            type="Research Paper"
        )
        db.add(new_doc)
        db.commit()
        db.refresh(new_doc)
        
        # Insert into ChromaDB
        import chromadb
        chroma_client = chromadb.PersistentClient(path="./chroma_db")
        collection = chroma_client.get_or_create_collection(name="climate_docs")
        
        text_to_embed = f"Title: {new_doc.title}\nSource: {new_doc.source}\nCountry: {new_doc.country}\nContent: {extracted_text[:3000]}"
        
        collection.upsert(
            documents=[text_to_embed],
            metadatas=[{
                "title": new_doc.title,
                "source": new_doc.source,
                "country": new_doc.country,
                "original_id": str(new_doc.id)
            }],
            ids=[f"doc_{new_doc.id}"]
        )

        return {
            "status": "success",
            "message": "File uploaded, parsed, and embedded successfully.",
            "original_name": file.filename,
            "saved_path": file_path
        }
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))
